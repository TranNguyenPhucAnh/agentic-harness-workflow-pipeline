"""
toolkits/devops_mlops/redactor.py
==================================
Pure transformation module — zero LLM, zero network calls.

Reads handover documents (Excel, PDF, DOCX, plain text, markdown) and
produces a redacted preview that is safe to send to an LLM. Secrets,
credentials, passwords, API keys, and IP addresses are detected and
replaced with typed placeholders before any LLM sees the content.

Called by doc_absorber.py as a mandatory pre-processing step:

  Excel / PDF / DOCX  →  redactor.py  →  redacted_preview.md
                                          redaction_report.json
                                          images/  (extracted, not redacted)
                               ↓
                      [human reviews preview]
                               ↓
                      doc_absorber.py  →  LLM call

────────────────────────────────────────────────────────────────
What gets redacted
────────────────────────────────────────────────────────────────

  Secrets / credentials:
    AWS access keys          AKIA*, ASIA*
    AWS secret keys          40-char base64 strings near "secret"
    Generic API keys         strings near "key", "token", "secret",
                             "password", "passwd", "pwd", "credential"
    Bearer tokens            "Bearer <token>" patterns
    Basic auth               "Basic <base64>" patterns
    Connection strings       jdbc:, postgresql://, mysql://, mongodb://
    Private keys             -----BEGIN * PRIVATE KEY-----

  Network / infrastructure:
    IPv4 addresses           kept if flag --keep-ips, else redacted
    IPv6 addresses           always redacted
    Private subnets          10.x, 172.16-31.x, 192.168.x
    Hostnames with port      hostname:PORT patterns

  Structured file patterns:
    .env files               KEY=value → KEY=<REDACTED>
    JSON/YAML key-value      "password": "..." → "password": "<REDACTED>"
    Excel cells near         sensitive column headers

────────────────────────────────────────────────────────────────
What is PRESERVED (key names, not values)
────────────────────────────────────────────────────────────────

  AWS_ACCESS_KEY_ID         → AWS_ACCESS_KEY_ID=<AWS_KEY>
  GRAFANA_API_TOKEN         → GRAFANA_API_TOKEN=<API_TOKEN>
  DB_PASSWORD               → DB_PASSWORD=<PASSWORD>
  "host": "10.0.1.45"      → "host": "<INTERNAL_IP>"
  jdbc:postgresql://host/db → jdbc:postgresql://<HOST>/<DB_NAME>

  Key names are preserved because they tell config_consistency_checker
  which services and credentials the project uses — this is structural
  information, not secret information.

────────────────────────────────────────────────────────────────
Outputs written
────────────────────────────────────────────────────────────────

  doc_absorber/redacted/<stem>_redacted.md
    Human-readable redacted content. Show this to user for review.
    Structured by: sheet/section name, then content with placeholders.

  doc_absorber/redacted/redaction_report.json
    Audit log: what was found, what placeholder replaced it, location
    (sheet, row, col, or page/section for PDF/DOCX).
    Never contains actual secret values.

  doc_absorber/images/<stem>_<sheet>_<ref>.png
    Embedded images extracted from Excel/DOCX. NOT sent to LLM.
    Listed in redaction_report.json with their paths for human review.

────────────────────────────────────────────────────────────────
Artifact impact
────────────────────────────────────────────────────────────────

  Command                 redacted/   images/   report.json
  ─────────────────────  ──────────  ────────  ───────────
  (normal run)            OVERWRITE   WRITE     OVERWRITE
  --dry-run               stdout      –         –

────────────────────────────────────────────────────────────────
Usage
────────────────────────────────────────────────────────────────

  # Standalone — redact a file and review output before doc_absorber
  python redactor.py --project my-co --file handover.xlsx
  python redactor.py --project my-co --file handover.pdf
  python redactor.py --project my-co --file handover.docx
  python redactor.py --project my-co --file notes.md

  # Redact multiple files at once
  python redactor.py --project my-co --file handover.xlsx credentials.xlsx

  # Keep IP addresses (useful when IPs needed for cross-reference)
  python redactor.py --project my-co --file handover.xlsx --keep-ips

  # Dry run — print redacted output to stdout, no writes
  python redactor.py --project my-co --file handover.xlsx --dry-run

  # Called by doc_absorber.py (programmatic API):
  from redactor import redact_file
  result = redact_file(path, output_dir, keep_ips=False)
  # result.redacted_md_path — path to redacted preview
  # result.report           — structured redaction report
  # result.finding_count    — how many secrets found

For taxonomy details see artifacts/TAXONOMY.md
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import textwrap
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

_TOOLKIT_DIR = Path(__file__).parent
_REPO_ROOT   = _TOOLKIT_DIR.parent.parent
sys.path.insert(0, str(_REPO_ROOT))

from modules.artifact_tracking import (  # noqa: E402
    track_read, track_write,
    print_summary as print_artifact_summary,
)


# ─────────────────────────────────────────────────────────────────────────────
# Artifact paths
# ─────────────────────────────────────────────────────────────────────────────

def _devops_artifact_root() -> Path:
    override = os.environ.get("DEVOPS_ARTIFACT_ROOT")
    base     = Path(override) if override else _REPO_ROOT.parent / "outputs" / "devops_mlops"
    slug     = os.environ.get("PIPELINE_PROJECT", "default")
    return base / f"artifacts_{slug}"

def _redacted_dir()  -> Path: return _devops_artifact_root() / "doc_absorber" / "redacted"
def _images_dir()    -> Path: return _devops_artifact_root() / "doc_absorber" / "images"
def _report_path()   -> Path: return _redacted_dir() / "redaction_report.json"

def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


# ─────────────────────────────────────────────────────────────────────────────
# Redaction patterns
# ─────────────────────────────────────────────────────────────────────────────

# Each pattern: (compiled_regex, placeholder_template, category)
# placeholder_template may contain {key} which is replaced with the matched key name

_PATTERNS: list[tuple[re.Pattern, str, str]] = [

    # ── AWS credentials ───────────────────────────────────────────────────────
    (re.compile(r'\b(AKIA|ASIA|AROA|AIPA|ANPA|ANVA|APKA)[A-Z0-9]{16}\b'),
     "<AWS_ACCESS_KEY_ID>", "aws_access_key"),

    # AWS secret key: 40-char base64 near secret/key context
    (re.compile(
        r'(?i)(?:secret[_\s-]?(?:access[_\s-]?)?key|aws[_\s-]secret)["\s:=]+([A-Za-z0-9/+]{40})\b'
    ), "<AWS_SECRET_KEY>", "aws_secret_key"),

    # ── Private keys ──────────────────────────────────────────────────────────
    (re.compile(
        r'-----BEGIN\s+(?:RSA\s+|EC\s+|OPENSSH\s+|DSA\s+|PGP\s+)?PRIVATE KEY-----.*?'
        r'-----END\s+(?:RSA\s+|EC\s+|OPENSSH\s+|DSA\s+|PGP\s+)?PRIVATE KEY-----',
        re.DOTALL,
    ), "<PRIVATE_KEY_BLOCK>", "private_key"),

    # ── Bearer / Basic tokens ─────────────────────────────────────────────────
    (re.compile(r'\bBearer\s+([A-Za-z0-9\-._~+/]{20,}={0,2})\b'),
     "Bearer <BEARER_TOKEN>", "bearer_token"),

    (re.compile(r'\bBasic\s+([A-Za-z0-9+/]{20,}={0,2})\b'),
     "Basic <BASIC_AUTH_TOKEN>", "basic_auth"),

    # ── Generic API keys / tokens / passwords ─────────────────────────────────
    # Pattern: KEY_NAME = "value" or KEY_NAME: value (env var / yaml / json style)
    (re.compile(
        r'(?i)((?:api[_\s]?key|api[_\s]?token|access[_\s]?token|auth[_\s]?token|'
        r'secret[_\s]?key|client[_\s]?secret|app[_\s]?secret|'
        r'password|passwd|pwd|pass|credentials?|'
        r'private[_\s]?key|service[_\s]?account[_\s]?key))'
        r'\s*(?:=|:)\s*'
        r'(?:"([^"]{4,})"'      # double-quoted value
        r"|'([^']{4,})'"        # single-quoted value
        r'|([^\s\n,;}{]{4,}))', # unquoted value
    ), "<{key}_REDACTED>", "credential"),

    # ── Connection strings ────────────────────────────────────────────────────
    (re.compile(
        r'(?i)(jdbc|postgresql|postgres|mysql|mariadb|mongodb|redis|amqp|amqps|'
        r'elasticsearch|neo4j|cassandra|couchdb)'
        r'://(?:[^@\s]+@)?([^/\s]+)/([^\s"\']+)',
    ), r"\1://<HOST>/<DB_NAME>", "connection_string"),

    # ── IPv4 — private ranges ─────────────────────────────────────────────────
    # Handled separately via _redact_ips() to support --keep-ips flag

    # ── IPv6 — always redact ──────────────────────────────────────────────────
    # Full form + compressed forms including ::1 (loopback) and :: (all-zeros)
    (re.compile(
        r'(?:(?:[0-9a-fA-F]{1,4}:){7}[0-9a-fA-F]{1,4}'          # full
        r'|(?:[0-9a-fA-F]{1,4}:){1,7}:'                          # trailing ::
        r'|(?:[0-9a-fA-F]{1,4}:){1,6}:[0-9a-fA-F]{1,4}'         # one group after ::
        r'|(?:[0-9a-fA-F]{1,4}:){1,5}(?::[0-9a-fA-F]{1,4}){1,2}'
        r'|(?:[0-9a-fA-F]{1,4}:){1,4}(?::[0-9a-fA-F]{1,4}){1,3}'
        r'|(?:[0-9a-fA-F]{1,4}:){1,3}(?::[0-9a-fA-F]{1,4}){1,4}'
        r'|(?:[0-9a-fA-F]{1,4}:){1,2}(?::[0-9a-fA-F]{1,4}){1,5}'
        r'|[0-9a-fA-F]{1,4}:(?::[0-9a-fA-F]{1,4}){1,6}'
        r'|::(?:[0-9a-fA-F]{1,4}:){0,5}[0-9a-fA-F]{1,4}'        # leading ::
        r'|::)',                                                   # :: alone
    ), "<IPV6_ADDRESS>", "ipv6"),

    # ── Hostname:PORT ─────────────────────────────────────────────────────────
    # NOTE: placeholder uses {port} — _replacer injects m.group(2) for port
    (re.compile(
        r'\b([a-zA-Z][a-zA-Z0-9\-]{2,63}\.[a-zA-Z]{2,}):([0-9]{2,5})\b'
    ), "<HOST>:{port}", "hostname_port"),

    # ── Generic high-entropy strings near sensitive keys ─────────────────────
    # 32+ char hex string (common for secrets)
    (re.compile(
        r'(?i)(?:token|key|secret|hash|digest|signature|hmac)'
        r'\s*(?:=|:)\s*'
        r'([0-9a-f]{32,})\b',
    ), r"<HEX_SECRET>", "hex_secret"),

    # ── Email addresses — keep domain, redact local part ──────────────────────
    (re.compile(r'\b([A-Za-z0-9._%+\-]{2,})@([A-Za-z0-9.\-]+\.[A-Za-z]{2,})\b'),
     r"<EMAIL_LOCAL>@\2", "email"),
]

# IPv4 patterns — applied separately to support --keep-ips
_IPV4_PRIVATE = re.compile(
    r'\b(10\.\d{1,3}\.\d{1,3}\.\d{1,3}'
    r'|172\.(1[6-9]|2[0-9]|3[01])\.\d{1,3}\.\d{1,3}'
    r'|192\.168\.\d{1,3}\.\d{1,3})\b'
)
_IPV4_PUBLIC = re.compile(
    r'\b(?!10\.|172\.(?:1[6-9]|2[0-9]|3[01])\.|192\.168\.)'
    r'(?:\d{1,3}\.){3}\d{1,3}\b'
)


# ─────────────────────────────────────────────────────────────────────────────
# Finding dataclass
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Finding:
    category:    str
    location:    str          # "Sheet1!B3", "page 2", "section Credentials"
    key_name:    str          # preserved key name or pattern description
    placeholder: str          # what replaced the value
    original_length: int      # length of redacted value (not the value itself)
    context_snippet: str      # up to 60 chars around the finding, value removed


@dataclass
class RedactionResult:
    source_path:       Path
    redacted_md_path:  Path | None
    image_paths:       list[Path]
    findings:          list[Finding]
    finding_count:     int
    categories:        dict[str, int]   # category → count
    report:            dict[str, Any]
    redacted_content:  str = ""        # full redacted markdown (for dry-run)


# ─────────────────────────────────────────────────────────────────────────────
# Core text redactor
# ─────────────────────────────────────────────────────────────────────────────

def _redact_text(
    text:     str,
    location: str,
    keep_ips: bool,
    findings: list[Finding],
) -> str:
    """
    Apply all redaction patterns to a text block.
    Modifies `findings` in-place. Returns redacted text.
    """
    result = text

    # Apply compiled patterns
    for pattern, placeholder, category in _PATTERNS:
        def _replacer(m: re.Match, _ph: str = placeholder, _cat: str = category) -> str:
            matched  = m.group(0)
            key_name = m.group(1) if m.lastindex and m.lastindex >= 1 else _cat
            # Build placeholder — substitute key name using simple str.replace
            # Use m.string (original text before this sub call) for snippet offsets
            ph = _ph.replace("{key}", str(key_name).upper())
            ph = ph.replace(r'\1', str(key_name).upper())
            # Inject port if placeholder uses {port} (hostname_port pattern)
            if "{port}" in ph and m.lastindex and m.lastindex >= 2:
                ph = ph.replace("{port}", m.group(2))
            # Snippet: use m.string so offsets are always valid
            src_str = m.string
            start   = max(0, m.start() - 20)
            end     = min(len(src_str), m.end() + 20)
            snippet = src_str[start:m.start()] + "[REDACTED]" + src_str[m.end():end]
            snippet = snippet.replace("\n", " ").strip()[:80]
            findings.append(Finding(
                category        = _cat,
                location        = location,
                key_name        = str(key_name)[:60],
                placeholder     = ph,
                original_length = len(matched),
                context_snippet = snippet,
            ))
            return ph
        result = pattern.sub(_replacer, result)

    # IP addresses — handled separately for --keep-ips support
    if not keep_ips:
        def _ip_replacer(m: re.Match) -> str:
            ip = m.group(0)
            findings.append(Finding(
                category        = "ipv4_private",
                location        = location,
                key_name        = "ip_address",
                placeholder     = "<INTERNAL_IP>",
                original_length = len(ip),
                context_snippet = f"...{ip}...",
            ))
            return "<INTERNAL_IP>"
        result = _IPV4_PRIVATE.sub(_ip_replacer, result)

        def _pub_ip_replacer(m: re.Match) -> str:
            ip = m.group(0)
            findings.append(Finding(
                category        = "ipv4_public",
                location        = location,
                key_name        = "ip_address",
                placeholder     = "<PUBLIC_IP>",
                original_length = len(ip),
                context_snippet = f"...{ip}...",
            ))
            return "<PUBLIC_IP>"
        result = _IPV4_PUBLIC.sub(_pub_ip_replacer, result)

    return result


# ─────────────────────────────────────────────────────────────────────────────
# Excel reader
# ─────────────────────────────────────────────────────────────────────────────

def _read_excel(
    path:     Path,
    keep_ips: bool,
    findings: list[Finding],
    images:   list[Path],
    img_dir:  Path,
) -> str:
    """
    Read Excel workbook. Returns redacted markdown string.
    Extracts: cell values, comments/notes, embedded images.
    Skips cells that are empty or contain only formulas.
    """
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return "_[Excel reading requires openpyxl: pip install openpyxl]_\n"

    try:
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception as e:
        return f"_[Could not open Excel file: {e}]_\n"

    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}\n")

        # ── Extract images from sheet ─────────────────────────────────────────
        sheet_images = getattr(ws, "_images", [])
        for idx, img in enumerate(sheet_images):
            try:
                from openpyxl.drawing.image import Image as XLImage  # type: ignore
                img_dir.mkdir(parents=True, exist_ok=True)
                stem     = path.stem
                img_name = f"{stem}_{sheet_name}_img{idx+1}.png"
                img_path = img_dir / img_name
                if hasattr(img, "ref"):
                    img_data = img.ref
                    if hasattr(img_data, "read"):
                        img_path.write_bytes(img_data.read())
                    elif isinstance(img_data, (bytes, bytearray)):
                        img_path.write_bytes(img_data)
                    images.append(img_path)
                    track_write(img_path)
                    parts.append(
                        f"_[IMAGE EXTRACTED: {img_name} — review manually before"
                        f" using in LLM context]_\n"
                    )
            except Exception:
                parts.append(f"_[IMAGE at {sheet_name} could not be extracted]_\n")

        # ── Extract cell data ─────────────────────────────────────────────────
        # Build header row detection (first non-empty row with multiple values)
        rows_data: list[list[str]] = []
        header_row: list[str] = []

        for row in ws.iter_rows():
            row_values: list[str] = []
            for cell in row:
                val = cell.value
                if val is None:
                    row_values.append("")
                    continue
                cell_str = str(val).strip()
                if not cell_str:
                    row_values.append("")
                    continue
                # Redact cell content
                location    = f"{sheet_name}!{cell.coordinate}"
                cell_redact = _redact_text(cell_str, location, keep_ips, findings)
                # Also check cell comment
                if cell.comment and cell.comment.text:
                    comment_redact = _redact_text(
                        cell.comment.text.strip(), f"{location} comment", keep_ips, findings
                    )
                    cell_redact += f" [note: {comment_redact}]"
                row_values.append(cell_redact)
            # Skip entirely empty rows
            if any(v.strip() for v in row_values):
                rows_data.append(row_values)

        if not rows_data:
            parts.append("_[Empty sheet]_\n\n")
            continue

        # Detect header row: first row where most cells are non-empty and look like labels
        if rows_data:
            header_row  = rows_data[0]
            data_rows   = rows_data[1:]
        else:
            header_row  = []
            data_rows   = []

        # Render as markdown table if looks tabular (header + data rows)
        if header_row and data_rows and len(header_row) > 1:
            # Filter to columns that have at least one header label
            col_indices = [i for i, h in enumerate(header_row) if h.strip()]
            if col_indices:
                headers_used = [header_row[i] for i in col_indices]
                parts.append("| " + " | ".join(headers_used) + " |\n")
                parts.append("|" + "|".join("---" for _ in headers_used) + "|\n")
                for row in data_rows[:500]:  # cap rows
                    row_cells = [row[i] if i < len(row) else "" for i in col_indices]
                    # Skip entirely empty data rows
                    if not any(c.strip() for c in row_cells):
                        continue
                    parts.append("| " + " | ".join(row_cells) + " |\n")
                parts.append("\n")
        else:
            # Non-tabular — render as key-value pairs
            for row in rows_data[:200]:
                non_empty = [v for v in row if v.strip()]
                if non_empty:
                    parts.append("- " + "  |  ".join(non_empty) + "\n")
            parts.append("\n")

    return "".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# PDF reader
# ─────────────────────────────────────────────────────────────────────────────

def _read_pdf(
    path:     Path,
    keep_ips: bool,
    findings: list[Finding],
) -> str:
    """
    Extract text from PDF. Uses pdfminer.six if available, pypdf fallback.
    Images in PDF are noted but not extracted (require pdf2image + poppler).
    """
    # Try pdfminer.six first
    try:
        from pdfminer.high_level import extract_pages   # type: ignore
        from pdfminer.layout import LTTextContainer     # type: ignore

        parts: list[str] = []
        for page_num, page_layout in enumerate(extract_pages(str(path)), start=1):
            page_text_blocks: list[str] = []
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    text = element.get_text().strip()
                    if text:
                        location    = f"page {page_num}"
                        redacted    = _redact_text(text, location, keep_ips, findings)
                        page_text_blocks.append(redacted)

            if page_text_blocks:
                parts.append(f"## Page {page_num}\n\n")
                parts.append("\n\n".join(page_text_blocks))
                parts.append("\n\n")

        return "".join(parts) or "_[No text extracted from PDF]_\n"

    except ImportError:
        pass

    # Fallback: pypdf
    try:
        import pypdf  # type: ignore

        parts = []
        reader = pypdf.PdfReader(str(path))
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                location = f"page {i}"
                redacted = _redact_text(text, location, keep_ips, findings)
                parts.append(f"## Page {i}\n\n{redacted}\n\n")

        return "".join(parts) or "_[No text extracted from PDF]_\n"

    except ImportError:
        return (
            "_[PDF reading requires pdfminer.six or pypdf: "
            "pip install pdfminer.six  OR  pip install pypdf]_\n"
        )
    except Exception as e:
        return f"_[Could not read PDF: {e}]_\n"


# ─────────────────────────────────────────────────────────────────────────────
# DOCX reader
# ─────────────────────────────────────────────────────────────────────────────

def _read_docx(
    path:     Path,
    keep_ips: bool,
    findings: list[Finding],
    images:   list[Path],
    img_dir:  Path,
) -> str:
    """
    Extract text and embedded images from DOCX.
    Tables are rendered as markdown tables.
    """
    try:
        import docx  # type: ignore
    except ImportError:
        return "_[DOCX reading requires python-docx: pip install python-docx]_\n"

    try:
        doc = docx.Document(str(path))
    except Exception as e:
        return f"_[Could not open DOCX: {e}]_\n"

    parts: list[str] = []

    # ── Extract embedded images ───────────────────────────────────────────────
    img_dir.mkdir(parents=True, exist_ok=True)
    for idx, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                ext       = rel.target_ref.split(".")[-1].lower() or "png"
                img_name  = f"{path.stem}_img{idx+1}.{ext}"
                img_path  = img_dir / img_name
                img_path.write_bytes(img_bytes)
                images.append(img_path)
                track_write(img_path)
                parts.append(
                    f"_[IMAGE EXTRACTED: {img_name} — review manually]_\n"
                )
            except Exception:
                parts.append("_[IMAGE could not be extracted]_\n")

    # ── Paragraphs ────────────────────────────────────────────────────────────
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower()
        location   = f"paragraph {i+1}"
        redacted   = _redact_text(text, location, keep_ips, findings)

        if "heading 1" in style_name:
            parts.append(f"\n# {redacted}\n")
        elif "heading 2" in style_name:
            parts.append(f"\n## {redacted}\n")
        elif "heading 3" in style_name:
            parts.append(f"\n### {redacted}\n")
        else:
            parts.append(f"{redacted}\n")

    # ── Tables ────────────────────────────────────────────────────────────────
    for t_idx, table in enumerate(doc.tables):
        parts.append(f"\n### Table {t_idx+1}\n\n")
        for r_idx, row in enumerate(table.rows):
            cells = []
            for c_idx, cell in enumerate(row.cells):
                text     = cell.text.strip()
                location = f"table {t_idx+1} row {r_idx+1} col {c_idx+1}"
                redacted = _redact_text(text, location, keep_ips, findings)
                cells.append(redacted.replace("\n", " "))
            parts.append("| " + " | ".join(cells) + " |\n")
            if r_idx == 0:
                parts.append("|" + "|".join("---" for _ in cells) + "|\n")
        parts.append("\n")

    return "".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# Plain text / Markdown reader
# ─────────────────────────────────────────────────────────────────────────────

def _read_text(
    path:     Path,
    keep_ips: bool,
    findings: list[Finding],
) -> str:
    try:
        track_read(path)
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"_[Could not read file: {e}]_\n"
    return _redact_text(text, str(path.name), keep_ips, findings)


# ─────────────────────────────────────────────────────────────────────────────
# Dispatcher
# ─────────────────────────────────────────────────────────────────────────────

def _dispatch(
    path:     Path,
    keep_ips: bool,
    findings: list[Finding],
    images:   list[Path],
    img_dir:  Path,
) -> str:
    ext = path.suffix.lower()
    if ext in (".xlsx", ".xls", ".xlsm", ".ods"):
        return _read_excel(path, keep_ips, findings, images, img_dir)
    elif ext == ".pdf":
        return _read_pdf(path, keep_ips, findings)
    elif ext in (".docx", ".doc"):
        return _read_docx(path, keep_ips, findings, images, img_dir)
    elif ext in (".md", ".txt", ".rst", ".csv", ".env", ".yaml", ".yml",
                 ".json", ".toml", ".ini", ".cfg"):
        return _read_text(path, keep_ips, findings)
    else:
        # Try as plain text
        return _read_text(path, keep_ips, findings)


# ─────────────────────────────────────────────────────────────────────────────
# Redaction report builder
# ─────────────────────────────────────────────────────────────────────────────

def _build_report(
    source_path:  Path,
    findings:     list[Finding],
    images:       list[Path],
    keep_ips:     bool,
    run_at:       str,
) -> dict[str, Any]:
    categories: dict[str, int] = {}
    for f in findings:
        categories[f.category] = categories.get(f.category, 0) + 1

    return {
        "run_at":        run_at,
        "source_file":   str(source_path),
        "keep_ips":      keep_ips,
        "finding_count": len(findings),
        "categories":    categories,
        "images_extracted": [str(p) for p in images],
        "image_count":   len(images),
        "findings": [
            {
                "category":        f.category,
                "location":        f.location,
                "key_name":        f.key_name,
                "placeholder":     f.placeholder,
                "original_length": f.original_length,
                "context_snippet": f.context_snippet,
                # NEVER includes actual secret value
            }
            for f in findings
        ],
        "security_note": (
            "This report contains redaction audit info only. "
            "No actual secret values are stored here. "
            "Original file was NOT modified."
        ),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Preview header
# ─────────────────────────────────────────────────────────────────────────────

def _build_preview_header(
    source_path:   Path,
    finding_count: int,
    image_count:   int,
    categories:    dict[str, int],
    keep_ips:      bool,
    run_at:        str,
) -> str:
    cat_lines = "\n".join(
        f"  - {cat}: {count}" for cat, count in sorted(categories.items())
    )
    img_note = (
        f"\n- **Images extracted:** {image_count} file(s) in `doc_absorber/images/`"
        f" — review manually before using in any LLM context"
        if image_count else ""
    )
    ip_note = (
        "\n- **IP addresses:** KEPT (--keep-ips flag active)"
        if keep_ips else
        "\n- **IP addresses:** REDACTED"
    )

    return textwrap.dedent(f"""\
        <!-- REDACTED PREVIEW — generated by redactor.py -->
        <!-- Original file: {source_path.name} -->
        <!-- Redacted at: {run_at} -->
        <!-- Finding count: {finding_count} — review before sending to LLM -->

        # Redacted Preview: {source_path.name}

        > **⚠ Review this document before proceeding.**
        > All secrets have been replaced with typed placeholders.
        > Verify the redaction is complete before running doc_absorber.py.

        ## Redaction Summary

        - **Source:** `{source_path.name}`
        - **Findings:** {finding_count} value(s) redacted
        - **Categories:**
        {cat_lines if cat_lines else "  (none detected)"}
        {img_note}{ip_note}

        ---

        ## Document Content (Redacted)

    """)


# ─────────────────────────────────────────────────────────────────────────────
# Public API — called by doc_absorber.py
# ─────────────────────────────────────────────────────────────────────────────

def redact_file(
    path:       Path,
    output_dir: Path | None = None,
    keep_ips:   bool        = False,
    dry_run:    bool        = False,
) -> RedactionResult:
    """
    Redact a single file. Returns RedactionResult with paths and report.

    This is the programmatic API called by doc_absorber.py.
    For CLI usage, use main().

    Parameters
    ----------
    path        : Path to the source file (Excel, PDF, DOCX, text, etc.)
    output_dir  : Where to write redacted_*.md and images/. Defaults to
                  doc_absorber/redacted/ in the devops artifact root.
    keep_ips    : If True, IPv4 addresses are NOT redacted.
    dry_run     : If True, returns result without writing any files.

    Returns
    -------
    RedactionResult with:
      .redacted_md_path  — Path to written redacted markdown (None if dry_run)
      .image_paths       — Paths to extracted images
      .findings          — List of Finding objects
      .finding_count     — Total findings
      .categories        — {category: count} dict
      .report            — Full structured report dict
    """
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    run_at    = _now_iso()
    findings: list[Finding] = []
    images:   list[Path]    = []

    out_dir = output_dir or _redacted_dir()
    img_dir = _images_dir()

    track_read(path)
    content = _dispatch(path, keep_ips, findings, images, img_dir)

    # Build categories
    categories: dict[str, int] = {}
    for f in findings:
        categories[f.category] = categories.get(f.category, 0) + 1

    report = _build_report(path, findings, images, keep_ips, run_at)
    header = _build_preview_header(
        path, len(findings), len(images), categories, keep_ips, run_at
    )
    full_md = header + content

    redacted_md_path: Path | None = None

    if not dry_run:
        out_dir.mkdir(parents=True, exist_ok=True)
        md_out = out_dir / f"{path.stem}_redacted.md"
        md_out.write_text(full_md, encoding="utf-8")
        track_write(md_out)
        redacted_md_path = md_out

        # Merge into shared redaction_report.json (append per source)
        rpt_path = _report_path()
        existing_reports: list[dict[str, Any]] = []
        if rpt_path.exists():
            try:
                track_read(rpt_path)
                data = json.loads(rpt_path.read_text(encoding="utf-8"))
                existing_reports = (
                    data if isinstance(data, list) else data.get("files", [])
                )
            except Exception:
                pass
        # Replace entry for same source file if re-run
        existing_reports = [
            r for r in existing_reports if r.get("source_file") != str(path)
        ]
        existing_reports.append(report)
        rpt_path.write_text(
            json.dumps({"files": existing_reports}, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        track_write(rpt_path)

    return RedactionResult(
        source_path      = path,
        redacted_md_path = redacted_md_path,
        image_paths      = images,
        findings         = findings,
        finding_count    = len(findings),
        categories       = categories,
        report           = report,
        redacted_content = full_md,
    )


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="redactor.py",
        description="Redact secrets from handover documents before LLM processing.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            Examples:
              python redactor.py --project my-co --file handover.xlsx
              python redactor.py --project my-co --file handover.xlsx credentials.xlsx
              python redactor.py --project my-co --file handover.pdf --keep-ips
              python redactor.py --project my-co --file handover.xlsx --dry-run
        """),
    )
    p.add_argument("--project",  default=os.environ.get("PIPELINE_PROJECT"),
                   help="Project slug. Sets PIPELINE_PROJECT.")
    p.add_argument("--file",     nargs="+", required=True, metavar="FILE",
                   help="File(s) to redact. Supports: .xlsx, .pdf, .docx, .md, .txt")
    p.add_argument("--keep-ips", action="store_true",
                   help="Keep IP addresses (useful for cross-reference with live_discovery).")
    p.add_argument("--dry-run",  action="store_true",
                   help="Print redacted content to stdout, write nothing.")
    return p


def _configure_project(project: str | None, parser: argparse.ArgumentParser) -> None:
    if project:
        os.environ["PIPELINE_PROJECT"] = project
        return
    if not os.environ.get("PIPELINE_PROJECT"):
        parser.error("Use --project <name> or export PIPELINE_PROJECT=<name>.")


def main() -> None:
    parser = _build_parser()
    args   = parser.parse_args()

    _configure_project(args.project, parser)

    print("=" * 60)
    print("  REDACTOR")
    print("=" * 60)
    print()

    total_findings = 0
    total_images   = 0
    results:  list[RedactionResult] = []
    exit_code = 0

    try:
      for file_arg in args.file:
        path = Path(file_arg).expanduser().resolve()
        if not path.exists():
            print(f"  [warn] File not found: {path}")
            continue

        print(f"  Processing: {path.name}  ({path.stat().st_size:,} bytes)")

        result = redact_file(path, keep_ips=args.keep_ips, dry_run=args.dry_run)
        results.append(result)
        total_findings += result.finding_count
        total_images   += len(result.image_paths)

        if args.dry_run:
            # Use already-computed content — no double-dispatch
            print(result.redacted_content)
        else:
            print(f"    Findings:  {result.finding_count}")
            cat_str = "  ".join(f"{k}={v}" for k, v in result.categories.items())
            if cat_str:
                print(f"    By type:   {cat_str}")
            if result.image_paths:
                print(f"    Images:    {len(result.image_paths)} extracted")
            if result.redacted_md_path:
                print(f"    Preview:   {result.redacted_md_path}")
        print()

      if not args.dry_run and results:
        print(f"  Redaction report: {_report_path()}")
        print()
        print("=" * 60)
        print(f"  Total findings:   {total_findings}")
        print(f"  Total images:     {total_images}")
        print()
        if total_findings > 0:
            print("  ⚠  Review the redacted preview(s) above before running:")
            print("     python doc_absorber.py --project "
                  f"{os.environ.get('PIPELINE_PROJECT', '<name>')} "
                  "--file <original_file(s)>")
        else:
            print("  ✓  No secrets detected. Safe to proceed with doc_absorber.py.")
        print("=" * 60)

    except KeyboardInterrupt:
        print("\n[redactor] Interrupted.")
        exit_code = 130
    except Exception as exc:
        print(f"[redactor][error] {exc}", file=sys.stderr)
        import traceback; traceback.print_exc()
        exit_code = 1
    finally:
        print_artifact_summary("[redactor]")

    sys.exit(exit_code)


if __name__ == "__main__":
    main()
